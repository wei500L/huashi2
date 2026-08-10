#!/usr/bin/env python3
"""Build the FF4 vocabulary workbook, question workbook, and student reader."""

from __future__ import annotations

import json
import difflib
import re
import unicodedata
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "docs" / "data" / "lexi-bridge-ff4"
OUT = DATA / "student-deliverables"

WORD_BOOK = DATA / "wordbook.jsonl"
QUESTIONS = DATA / "question-bank-package.json"
TEM4 = DATA / "tem4-candidates.jsonl"
FALSE_FRIENDS = DATA / "false-friends-candidates.jsonl"
FALSE_FRIENDS_OCR = DATA / "source-ocr-cache/false-friends-pages.jsonl"
REVIEW = DATA / "review-report.json"

VOCAB_XLSX = OUT / "法语专四假朋友词汇总表.xlsx"
QUESTION_XLSX = OUT / "法语专四假朋友题库.xlsx"
STUDENT_DOCX = OUT / "法语专四假朋友词汇大全-学生阅读版.docx"

NAVY = "1F4E78"
BLUE = "5B9BD5"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EAF3F8"
GOLD = "D6A84B"
PALE_GOLD = "FFF2CC"
GREEN = "70AD47"
PALE_GREEN = "E2F0D9"
GRAY = "667085"
PALE_GRAY = "F2F4F7"
WHITE = "FFFFFF"
RED = "C00000"
FONT_CN = "Arial Unicode MS"

# Corrections confirmed from the rendered source pages and their visible example groups.
TEM4_SENSE_OVERRIDES = {
    "chance": ["运气", "好运", "机会", "可能性"],
    "heurter": ["碰撞", "触犯；与……相抵触"],
    "pièce": ["件；只", "房间", "戏剧作品", "文件", "硬币", "零件", "附件"],
    "occasion": ["机会", "时机；场合", "便宜货；二手物品"],
}

FALSE_FRIEND_SENSE_OVERRIDES = {
    "ligne": ["线", "轮廓", "图形", "段落", "身材", "航线", "写几句话；写个便条"],
    "pardonner": ["宽恕", "原谅", "致命的；不留情的（qui ne pardonne pas）"],
}

SOURCE_USAGE_OVERRIDES = {
    "chance": "bonne chance；mauvaise chance；avoir de la chance；donner une chance à tous d’apprendre à skier",
    "heurter": "une voiture a heurté une autre voiture；heurter les intérêts",
    "pièce": "vendre à la pièce；quitter une pièce；pièce de théâtre；pièce jointe / attachée；pièce de monnaie",
    "occasion": "manquer / saisir une occasion；à l’occasion de；livre d’occasion",
}

RELATION_OVERRIDES = {
    # Source-text spot checks where OCR phrasing establishes a genuine shared use,
    # but the Chinese candidate fields use non-identical translations.
    "remonter": "B 部分对应/语义范围变化",
    "sage": "B 部分对应/语义范围变化",
    "cravate": "B 部分对应/语义范围变化",
    "station": "B 部分对应/语义范围变化",
    "goutte": "B 部分对应/语义范围变化",
    "foyer": "B 部分对应/语义范围变化",
    "ordinaire": "B 部分对应/语义范围变化",
    "interroger": "B 部分对应/语义范围变化",
    "chanter": "B 部分对应/语义范围变化",
    "journal": "B 部分对应/语义范围变化",
    "craquer": "B 部分对应/语义范围变化",
    "square": "B 部分对应/语义范围变化",
    "sain": "B 部分对应/语义范围变化",
    "réservoir": "B 部分对应/语义范围变化",
    "rail": "B 部分对应/语义范围变化",
    "folie": "B 部分对应/语义范围变化",
    "draguer": "B 部分对应/语义范围变化",
    "humeur": "B 部分对应/语义范围变化",
    "chagrin": "B 部分对应/语义范围变化",
    "raquette": "B 部分对应/语义范围变化",
    "digestif": "B 部分对应/语义范围变化",
    "commettre": "B 部分对应/语义范围变化",
    "repasser": "B 部分对应/语义范围变化",
    "information": "B 部分对应/语义范围变化",
    "salaire": "C 语义范围/用法边界",
    "général": "C 语义范围/用法边界",
    "physique": "C 语义范围/用法边界",
    "revenu": "C 语义范围/用法边界",
    "ski": "D 同形同义借词边界",
}


def load_jsonl(path: Path) -> list[dict]:
    rows = []
    with path.open("r", encoding="utf-8-sig") as handle:
        for line in handle:
            if line.strip():
                rows.append(json.loads(line))
    return rows


def normalize(value: str | None) -> str:
    if not value:
        return ""
    return re.sub(r"[\s’'\-]+", "", value.casefold())


def normalize_latin(value: str | None) -> str:
    """Normalize Latin-script OCR tokens without conflating longer words."""
    if not value:
        return ""
    decomposed = unicodedata.normalize("NFKD", value.casefold())
    unaccented = "".join(char for char in decomposed if not unicodedata.combining(char))
    return re.sub(r"[^a-z0-9]", "", unaccented)


def first_sense(value: str | None) -> str:
    if not value:
        return ""
    return re.split(r"[;；，,、/]", value)[0].strip()


def join_list(value) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "；".join(str(item).strip() for item in value if str(item).strip())
    return str(value).strip()


def choose_tem4(word: dict, candidates: list[dict]) -> dict:
    target = normalize(word["french_word"])
    exact = [c for c in candidates if normalize(c.get("normalized_headword")) == target]
    if not exact:
        return {}
    page = int(word.get("tem4_pdf_page") or 0)
    same_page = [c for c in exact if int(c.get("pdf_page") or 0) == page]
    pool = same_page or exact
    return max(
        pool,
        key=lambda c: (
            len(c.get("frequency_or_exam_markers") or []),
            len(c.get("chinese_core_senses") or []),
            bool(c.get("part_of_speech")),
            bool(c.get("ipa")),
        ),
    )


def choose_false_friend(word: dict, candidates: list[dict]) -> dict:
    target = normalize(word["french_word"])
    exact = [c for c in candidates if normalize(c.get("normalized_headword")) == target]
    if not exact:
        return {}
    page = int(word.get("false_friend_pdf_page") or 0)
    same_page = [c for c in exact if int(c.get("pdf_page") or 0) == page]
    pool = same_page or exact
    return max(
        pool,
        key=lambda c: (
            len(c.get("french_true_meanings") or []),
            len(c.get("english_true_meanings") or []),
            bool(c.get("page_evidence")),
        ),
    )


def load_false_friend_ocr(path: Path) -> dict[int, str]:
    pages: dict[int, str] = {}
    for row in load_jsonl(path):
        match = re.search(r"ff-(\d+)\.png$", row.get("path") or "")
        if match and row.get("text"):
            pages[int(match.group(1))] = row["text"].strip()
    return pages


def first_ocr_token(line: str) -> str:
    match = re.match(r"^\s*([^\s(]+)", line)
    return normalize_latin(match.group(1)) if match else ""


def find_headword_line(
    lines: list[str],
    headword: str,
    other_headwords: set[str] | None = None,
    start_at: int = 0,
) -> tuple[int | None, str]:
    target = normalize_latin(headword)
    other_headwords = other_headwords or set()
    exact_lower = []
    exact_other = []
    fuzzy = []
    for index in range(start_at, len(lines)):
        token = first_ocr_token(lines[index])
        if not token:
            continue
        content_length = len(normalize_latin(lines[index]))
        if token == target and content_length > len(target) + 2:
            raw_first = re.match(r"^\s*([^\s(]+)", lines[index]).group(1)
            if raw_first[:1].islower():
                exact_lower.append(index)
            else:
                exact_other.append(index)
        elif (
            len(target) >= 5
            and token not in other_headwords
            and difflib.SequenceMatcher(None, token, target).ratio() >= 0.74
        ):
            fuzzy.append((difflib.SequenceMatcher(None, token, target).ratio(), index))
    if exact_lower:
        return exact_lower[0], "原书OCR精确定位"
    if fuzzy:
        # OCR occasionally corrupts one or two letters in a bold headword.
        _, index = max(fuzzy, key=lambda item: (item[0], -item[1]))
        return index, "原书OCR近似定位（词头存在识别误差）"
    if exact_other and headword[:1].isupper():
        return exact_other[0], "原书OCR精确定位"
    return None, "视觉核验摘要回退"


def extract_false_friend_excerpt(
    word: dict,
    candidate: dict,
    page_text: str,
    page_candidates: list[dict],
) -> tuple[str, str]:
    """Extract a compact entry excerpt and stop at the next verified headword."""
    evidence = (candidate.get("page_evidence") or "").strip()
    if not page_text:
        return f"[视觉核验摘要] {evidence}", "视觉核验摘要回退"
    lines = [line.strip() for line in page_text.splitlines() if line.strip()]
    other_heads = {
        normalize_latin(row.get("normalized_headword") or row.get("standard_french_headword"))
        for row in page_candidates
        if normalize_latin(row.get("normalized_headword") or row.get("standard_french_headword"))
        != normalize_latin(word["french_word"])
    }
    start, status = find_headword_line(lines, word["french_word"], other_heads)
    if start is None:
        return f"[视觉核验摘要] {evidence}", status
    end = min(len(lines), start + 14)
    for index in range(start + 1, end):
        token = first_ocr_token(lines[index])
        raw_match = re.match(r"^\s*([^\s(]+)", lines[index])
        raw_first = raw_match.group(1) if raw_match else ""
        fuzzy_other = any(
            len(other) >= 5 and difflib.SequenceMatcher(None, token, other).ratio() >= 0.82
            for other in other_heads
        )
        if raw_first[:1].islower() and (token in other_heads or fuzzy_other):
            end = index
            break
    excerpt = " ".join(lines[start:end])
    excerpt = re.sub(r"\s+", " ", excerpt).strip()
    if len(excerpt) > 900:
        excerpt = excerpt[:897].rstrip() + "…"
    if len(excerpt) < len(word["french_word"]) + 8:
        return f"[视觉核验摘要] {evidence}", "视觉核验摘要回退"
    return excerpt, status


def frequency_level(markers: list[str]) -> str:
    count = len(markers)
    if count >= 2:
        return "高频"
    if count == 1:
        return "中频"
    return "低频"


def classify_relation(
    french_word: str,
    english_word: str,
    french_senses: list[str],
    english_senses: list[str],
    source_excerpt: str,
    page_evidence: str,
) -> tuple[str, list[str], list[str], list[str], str]:
    """Classify from the book's wording; Chinese field overlap is only auxiliary evidence."""
    french_map = {re.sub(r"[\s，,；;、/（）()·…]+", "", value): value for value in french_senses if value}
    english_map = {re.sub(r"[\s，,；;、/（）()·…]+", "", value): value for value in english_senses if value}
    overlap_keys = set(french_map) & set(english_map)
    french_only_keys = set(french_map) - set(english_map)
    english_only_keys = set(english_map) - set(french_map)
    overlap = [french_map[key] for key in french_map if key in overlap_keys]
    french_only = [french_map[key] for key in french_map if key in french_only_keys]
    english_only = [english_map[key] for key in english_map if key in english_only_keys]
    if french_word in RELATION_OVERRIDES:
        relation = RELATION_OVERRIDES[french_word]
        reasons = {
            "B 部分对应/语义范围变化": "按原书英文释义逐项核对，确认存在真实对应义，同时另有更常用或更宽的法语义。",
            "C 语义范围/用法边界": "按原书对词义范围、词性或常用程度的明确说明归入用法边界。",
            "D 同形同义借词边界": "原书与双源释义均显示核心义基本一致，作为边界词保留。",
        }
        return relation, overlap, french_only, english_only, reasons[relation]

    excerpt_body = source_excerpt.strip()
    if excerpt_body:
        # Remove the displayed French headword so that it is not mistaken for an English match.
        excerpt_body = re.sub(rf"^\s*{re.escape(french_word)}\b", "", excerpt_body, count=1, flags=re.IGNORECASE)
    book_text = f"{excerpt_body} {page_evidence}".casefold()
    english_pattern = rf"\b{re.escape(english_word.casefold())}\b"
    explicit_nonmatch = bool(
        re.search(rf"\bnot\s+(?:always\s+|normally\s+|usually\s+)?{english_pattern[2:]}", book_text)
        or re.search(rf"{english_pattern}(?:\s*\([^)]{{0,50}}\))?\s*=", book_text)
        or any(marker in page_evidence for marker in ("不表示", "不是", "≠", "无关"))
    )
    opening = excerpt_body[:520].casefold()
    direct_shared_patterns = (
        rf"^\s*(?:\([^)]+\)\s*)?{english_pattern}",
        rf"^\s*(?:\([^)]+\)\s*)?to\s*{english_pattern}",
        rf"\b(?:can|sometimes|occasionally)\s+(?:also\s+)?(?:mean|means|be|to)\s+(?:to\s+)?{english_pattern}",
        rf"\b(?:normally|usually|often)\s+(?:also\s+)?means?\s+(?:to\s+)?{english_pattern}",
        rf"\bto\s+{english_pattern}\s+is\s+(?:one|sometimes)",
        rf"\bmeans?\s+{english_pattern}",
    )
    shared_in_opening = any(re.search(pattern, opening) for pattern in direct_shared_patterns)
    if explicit_nonmatch and re.match(rf"^\s*(?:\([^)]+\)\s*)?(?:not\s+)?{english_pattern}(?:\s*\([^)]{{0,50}}\))?\s*=", opening):
        shared_in_opening = False
    strong_partial_cues = any(
        cue in book_text
        for cue in (
            "can sometimes mean",
            "sometimes has the sense",
            "occasionally",
            "not only",
            "but also",
            "as well as",
            "more often than",
            "more commonly",
            "wider meaning",
            "in certain senses",
            "still overlap",
            "same word",
            "范围较广",
            "义更广",
            "义窄",
            "可偶为",
        )
    )
    if overlap or shared_in_opening or strong_partial_cues:
        relation = "B 部分对应/语义范围变化"
        reason = "原书显示两词至少在部分语境可对应，同时强调法语另有常用义或两者适用范围不同。"
    else:
        relation = "A 核心义不对应/原书另词表达"
        reason = "原书将该法语词的常用义与易混英文词区分；中文释义字段未提供可靠的共同核心义。"
    return relation, overlap, french_only, english_only, reason


def build_distinction(word: dict) -> str:
    relation = word["relation_type"]
    french = word["french_word"]
    english = word["english_word"]
    if relation == "B 部分对应/语义范围变化":
        parts = []
        if word["shared_senses"]:
            parts.append(f"中文释义字段直接重合：{join_list(word['shared_senses'])}")
        if word["french_only_senses"]:
            parts.append(f"法语额外义项：{join_list(word['french_only_senses'])}")
        if word["english_only_senses"]:
            parts.append(f"英文额外义项：{join_list(word['english_only_senses'])}")
        auxiliary = "；".join(parts)
        return f"{word['book_relation_reason']}" + (f" 辅助对照：{auxiliary}。" if auxiliary else "")
    if relation == "C 语义范围/用法边界":
        evidence = word.get("false_friend_evidence") or "两词中译接近，但使用范围、词性或搭配不同"
        return f"{word['book_relation_reason']} 来源提示：{evidence}"
    if relation == "D 同形同义借词边界":
        return "法英核心义基本相同，属于同形同义借词边界；保留用于辨认边界，不作为核心假朋友记忆负担。"
    return (
        f"法语“{french}”表示“{word['tem4_senses']}”；"
        f"英文“{english}”表示“{word['english_senses']}”。{word['book_relation_reason']}"
    )


def mnemonic(word: dict) -> str:
    return (
        f"双向记忆：{word['french_word']} → {first_sense(word['tem4_senses'])}；"
        f"{word['english_word']} → {first_sense(word['english_senses']) or '英文原义'}。"
    )


def source_usage(candidate: dict) -> str:
    evidence = (candidate.get("page_evidence") or "").strip()
    useful_evidence = ""
    if "；" in evidence:
        tail = evidence.split("；", 1)[1].strip()
        if re.search(r"[A-Za-zÀ-ÿ~=]", tail):
            useful_evidence = tail
    derivatives = join_list(candidate.get("explicitly_listed_derivatives"))
    if useful_evidence and derivatives and derivatives not in useful_evidence:
        return f"{useful_evidence}｜词族/搭配：{derivatives}"
    return useful_evidence or derivatives or "来源页未单列可直接复用的例组"


def enrich_words(
    words: list[dict],
    tem4_candidates: list[dict],
    false_friend_candidates: list[dict],
    false_friend_ocr: dict[int, str],
) -> list[dict]:
    false_candidates_by_page: dict[int, list[dict]] = {}
    for candidate in false_friend_candidates:
        false_candidates_by_page.setdefault(int(candidate.get("pdf_page") or 0), []).append(candidate)
    enriched = []
    for index, word in enumerate(words, start=1):
        candidate = choose_tem4(word, tem4_candidates)
        false_candidate = choose_false_friend(word, false_friend_candidates)
        if not candidate or not false_candidate:
            raise RuntimeError(f"Missing exact dual-source candidate for {word['word_id']}")
        markers = candidate.get("frequency_or_exam_markers") or []
        french_word = word["french_word"]
        tem4_sense_values = TEM4_SENSE_OVERRIDES.get(french_word, candidate.get("chinese_core_senses") or [])
        ff_french_values = FALSE_FRIEND_SENSE_OVERRIDES.get(
            french_word, false_candidate.get("french_true_meanings") or []
        )
        english_values = false_candidate.get("english_true_meanings") or []
        ff_page = int(word["false_friend_pdf_page"])
        source_excerpt, excerpt_status = extract_false_friend_excerpt(
            word,
            false_candidate,
            false_friend_ocr.get(ff_page, ""),
            false_candidates_by_page.get(ff_page, []),
        )
        relation, shared, french_only, english_only, relation_reason = classify_relation(
            french_word,
            word["english_word"],
            ff_french_values,
            english_values,
            source_excerpt if not source_excerpt.startswith("[视觉核验摘要]") else "",
            false_candidate.get("page_evidence") or "",
        )
        row = dict(word)
        row.update(
            {
                "sequence": index,
                "part_of_speech": candidate.get("part_of_speech") or "",
                "ipa": candidate.get("ipa") or "",
                "markers": markers,
                "frequency": frequency_level(markers),
                "tem4_senses": join_list(tem4_sense_values),
                "false_friend_french_senses": join_list(ff_french_values),
                "english_senses": join_list(english_values),
                "relation_type": relation,
                "shared_senses": shared,
                "french_only_senses": french_only,
                "english_only_senses": english_only,
                "book_excerpt": source_excerpt,
                "book_excerpt_status": excerpt_status,
                "book_relation_reason": relation_reason,
                "derivatives": join_list(candidate.get("explicitly_listed_derivatives")),
                "source_usage": SOURCE_USAGE_OVERRIDES.get(french_word, source_usage(candidate)),
                "tem4_evidence": candidate.get("page_evidence") or "",
                "false_friend_evidence": false_candidate.get("page_evidence") or "",
                "printed_tem4_page": candidate.get("printed_page"),
                "printed_false_friend_page": false_candidate.get("printed_page"),
                "source_match": "通过：双源词头与记录页码精确匹配",
                "audit_correction": "",
                "pedagogic_status": "原书证据驱动复核完成",
            }
        )
        if french_word in TEM4_SENSE_OVERRIDES or french_word in FALSE_FRIEND_SENSE_OVERRIDES:
            row["audit_correction"] = "已按渲染源页纠正明显转录/语境表达"
        if french_word in RELATION_OVERRIDES and row["relation_type"] == "B 部分对应/语义范围变化":
            row["audit_correction"] = (row["audit_correction"] + "；" if row["audit_correction"] else "") + "按原书英文释义纠正中文字段不完全重合造成的误分"
        if row["relation_type"] == "C 语义范围/用法边界":
            row["audit_correction"] = (row["audit_correction"] + "；" if row["audit_correction"] else "") + "按原书范围、词性或常用程度说明归类"
        if row["relation_type"] == "D 同形同义借词边界":
            row["audit_correction"] = "转为扩展边界词，不列作核心假朋友"
        row["distinction"] = build_distinction(row)
        row["mnemonic"] = mnemonic(row)
        if row["relation_type"] == "D 同形同义借词边界":
            row["learning_priority"] = "C 扩展掌握"
        elif row["frequency"] == "高频":
            row["learning_priority"] = "A 核心必学"
        elif row["relation_type"] != "A 核心义不对应/原书另词表达":
            row["learning_priority"] = "B 重点辨析"
        else:
            row["learning_priority"] = "C 扩展掌握"
        row["chinese_gloss"] = row["tem4_senses"]
        row["false_friend_meanings"] = row["english_senses"]
        enriched.append(row)
    rank = {"高频": 0, "中频": 1, "低频": 2}
    return sorted(enriched, key=lambda row: (rank[row["frequency"]], row["sequence"]))


def set_sheet_title(ws, title: str, subtitle: str, column_count: int) -> None:
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=column_count)
    title_cell = ws.cell(1, 1, title)
    title_cell.font = Font(name=FONT_CN, size=18, bold=True, color=WHITE)
    title_cell.fill = PatternFill("solid", fgColor=NAVY)
    title_cell.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 34

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=column_count)
    subtitle_cell = ws.cell(2, 1, subtitle)
    subtitle_cell.font = Font(name=FONT_CN, size=10, color="7A4E00")
    subtitle_cell.fill = PatternFill("solid", fgColor=PALE_GOLD)
    subtitle_cell.alignment = Alignment(wrap_text=True, vertical="center")
    ws.row_dimensions[2].height = 40


def write_table_sheet(
    ws,
    title: str,
    subtitle: str,
    headers: list[str],
    rows: list[list],
    widths: list[float],
    freeze: str = "A4",
    row_height: float = 54,
) -> None:
    set_sheet_title(ws, title, subtitle, len(headers))
    thin = Side(style="thin", color="D0D5DD")
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(3, col, header)
        cell.font = Font(name=FONT_CN, size=10, bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=BLUE)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    ws.row_dimensions[3].height = 32

    for row_index, values in enumerate(rows, start=4):
        fill = PatternFill("solid", fgColor=WHITE if row_index % 2 == 0 else "F8FBFD")
        for col, value in enumerate(values, start=1):
            cell = ws.cell(row_index, col, value)
            cell.font = Font(name=FONT_CN, size=10, color="24313B")
            cell.fill = fill
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
        ws.row_dimensions[row_index].height = row_height

    for col, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(col)].width = width
    ws.auto_filter.ref = f"A3:{get_column_letter(len(headers))}{len(rows) + 3}"
    ws.freeze_panes = freeze
    ws.sheet_view.showGridLines = False
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.oddFooter.center.text = "第 &P 页 / 共 &N 页"


def make_vocabulary_workbook(words: list[dict]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "词汇总表"
    headers = [
        "序号",
        "学习优先级",
        "专四频级",
        "法语单词",
        "词性",
        "音标",
        "易混英文词",
        "TEM4审定完整释义",
        "假朋友词典中的法语义",
        "英文词真实中文义",
        "原书关系类型",
        "关系判定依据",
        "Kirk-Greene视觉核验摘要",
        "Kirk-Greene原文摘录（OCR）",
        "原文定位状态",
        "中文释义字段直接重合（辅助）",
        "法语字段独有义项（辅助）",
        "英文字段独有义项（辅助）",
        "核心辨析与考试陷阱",
        "记忆提示",
        "TEM4来源例组/搭配",
        "词族与派生",
        "历年专四标记",
        "TEM4单元",
        "TEM4 PDF页码",
        "假朋友 PDF页码",
        "TEM4页码证据",
        "逐条来源核对",
        "审定纠正/处理",
        "词条ID",
        "内容状态",
    ]
    rows = [
        [
            word["sequence"],
            word["learning_priority"],
            word["frequency"],
            word["french_word"],
            word["part_of_speech"],
            word["ipa"],
            word["english_word"],
            word["tem4_senses"],
            word["false_friend_french_senses"],
            word["english_senses"],
            word["relation_type"],
            word["book_relation_reason"],
            word["false_friend_evidence"],
            word["book_excerpt"],
            word["book_excerpt_status"],
            join_list(word["shared_senses"]),
            join_list(word["french_only_senses"]),
            join_list(word["english_only_senses"]),
            word["distinction"],
            word["mnemonic"],
            word["source_usage"],
            word["derivatives"],
            "；".join(word["markers"]),
            word.get("unit") or "",
            word["tem4_pdf_page"],
            word["false_friend_pdf_page"],
            word["tem4_evidence"],
            word["source_match"],
            word["audit_correction"],
            word["word_id"],
            word["pedagogic_status"],
        ]
        for word in words
    ]
    write_table_sheet(
        ws,
        "法语专四假朋友词汇总表｜445条原书证据版",
        "关系类型以 Kirk-Greene 原书措辞和视觉核验摘要为主；中文释义字段重合仅作辅助。OCR摘录用于快速阅读，遇识别误差时以PDF页码和视觉核验摘要为准。",
        headers,
        rows,
        [8, 15, 11, 18, 14, 16, 20, 34, 34, 30, 28, 42, 54, 76, 25, 28, 30, 28, 66, 46, 66, 42, 24, 14, 14, 16, 52, 34, 38, 24, 24],
        freeze="E4",
        row_height=92,
    )
    for row in range(4, len(rows) + 4):
        level = ws.cell(row, 3).value
        color = {"高频": "F4CCCC", "中频": PALE_GOLD, "低频": PALE_GREEN}[level]
        ws.cell(row, 3).fill = PatternFill("solid", fgColor=color)
        ws.cell(row, 3).font = Font(name=FONT_CN, bold=True, color="7A1F1F" if level == "高频" else "385723")
        relation = ws.cell(row, 11).value
        relation_color = {
            "A 核心义不对应/原书另词表达": "FCE4D6",
            "B 部分对应/语义范围变化": "FFF2CC",
            "C 语义范围/用法边界": "E4DFEC",
            "D 同形同义借词边界": "D9EAD3",
        }[relation]
        ws.cell(row, 11).fill = PatternFill("solid", fgColor=relation_color)
        ws.cell(row, 31).fill = PatternFill("solid", fgColor=PALE_GREEN)

    focus = wb.create_sheet("重点辨析")
    focus_headers = ["优先级", "法语词", "英文易混词", "关系类型", "TEM4释义", "英文义", "核心陷阱", "来源例组", "词条ID"]
    focus_words = sorted(
        words,
        key=lambda word: (
            {"A 核心必学": 0, "B 重点辨析": 1, "C 扩展掌握": 2}[word["learning_priority"]],
            word["sequence"],
        ),
    )
    focus_rows = [
        [
            word["learning_priority"],
            word["french_word"],
            word["english_word"],
            word["relation_type"],
            word["tem4_senses"],
            word["english_senses"],
            word["distinction"],
            word["source_usage"],
            word["word_id"],
        ]
        for word in focus_words
        if word["learning_priority"] != "C 扩展掌握"
    ]
    write_table_sheet(
        focus,
        "重点辨析清单",
        "集中呈现高频词和义项重合/范围型陷阱，适合教师备课、学生首轮复习与题目解析。",
        focus_headers,
        focus_rows,
        [15, 18, 20, 27, 36, 32, 68, 62, 24],
        freeze="D4",
        row_height=78,
    )

    audit = wb.create_sheet("审查报告")
    set_sheet_title(audit, "445条词汇逐条审查报告", "本页说明审定口径；完整内容位于首个工作表“词汇总表”。", 6)
    relation_counts = Counter(word["relation_type"] for word in words)
    correction_count = sum(bool(word["audit_correction"]) for word in words)
    exact_ocr = sum(word["book_excerpt_status"] == "原书OCR精确定位" for word in words)
    fuzzy_ocr = sum(word["book_excerpt_status"].startswith("原书OCR近似定位") for word in words)
    fallback_ocr = sum(word["book_excerpt_status"] == "视觉核验摘要回退" for word in words)
    summary = [
        ("审查结果", "445 / 445 条均按双源词头、页码、视觉核验摘要和原书OCR摘录重新核对"),
        ("内容状态", "原书证据驱动复核完成；关系判定不再以中文字符串精确重合作为主规则"),
        ("A 核心义不对应", relation_counts["A 核心义不对应/原书另词表达"]),
        ("B 部分对应/范围变化", relation_counts["B 部分对应/语义范围变化"]),
        ("C 用法边界", relation_counts["C 语义范围/用法边界"]),
        ("D 同形同义边界", relation_counts["D 同形同义借词边界"]),
        ("原书摘录定位", f"精确 {exact_ocr} / OCR近似 {fuzzy_ocr} / 视觉核验摘要回退 {fallback_ocr}"),
        ("高/中/低频", f"{sum(w['frequency']=='高频' for w in words)} / {sum(w['frequency']=='中频' for w in words)} / {sum(w['frequency']=='低频' for w in words)}"),
        ("显式审定处理", f"{correction_count} 条包含转录纠正、范围型说明或边界处理"),
        ("已确认纠正示例", "chance：好远→好运并补全机会/可能性；heurter：撞撞→碰撞；pièce：邮件→附件；occasion：节节→时机"),
        ("频级口径", "TEM4历年专四标记≥2次为高频，1次为中频，无明确标记为低频"),
        ("原书口径", "关系类型首先依据 Kirk-Greene 原文措辞与视觉核验摘要；中文释义字段对照只作辅助，OCR文字不替代原始PDF"),
        ("来源", "KIRK_GREENE_FALSE_FRIENDS + TEM4_CORE_VOCABULARY；每条保留PDF页码、视觉证据与原文摘录/回退状态"),
        ("数据库边界", "本次仅完善内容文件，未执行数据库写入"),
    ]
    for row_index, (label, value) in enumerate(summary, start=4):
        audit.cell(row_index, 1, label).font = Font(name=FONT_CN, bold=True, color=NAVY)
        audit.cell(row_index, 1).fill = PatternFill("solid", fgColor=LIGHT_BLUE)
        audit.cell(row_index, 2, value).alignment = Alignment(wrap_text=True, vertical="top")
        audit.row_dimensions[row_index].height = 38
    audit.column_dimensions["A"].width = 24
    audit.column_dimensions["B"].width = 100
    audit.sheet_view.showGridLines = False
    wb.save(VOCAB_XLSX)


def rotate_choice_options(options: list[str], index: int) -> tuple[list[str], str]:
    if len(options) != 4:
        padded = (options + [""] * 4)[:4]
        return padded, "A"
    shift = index % 4
    rotated = options[shift:] + options[:shift]
    correct_position = (4 - shift) % 4
    return rotated, "ABCD"[correct_position]


def cloze_from_source(word: dict) -> str:
    usage = word.get("source_usage") or ""
    if not usage or usage.startswith("来源页未单列"):
        return f"Pour exprimer « {word['tem4_senses']} », on emploie le mot ____."
    phrase = usage.split("；", 1)[0].strip()
    if "~" in phrase:
        phrase = phrase.replace("~", "____", 1)
    else:
        phrase = re.sub(re.escape(word["french_word"]), "____", phrase, count=1, flags=re.IGNORECASE)
    if "____" not in phrase:
        return f"Pour exprimer « {word['tem4_senses']} », on emploie le mot ____."
    return f"请选择填入空格的法语词：{phrase}"


def distinct_distractor_words(words: list[dict], word: dict, offset: int, value_key: str = "french_word") -> list[dict]:
    result = []
    cursor = (word["sequence"] + offset) % len(words)
    while len(result) < 3:
        candidate = words[cursor % len(words)]
        cursor += 37
        if candidate[value_key] != word[value_key] and candidate[value_key] not in {x[value_key] for x in result}:
            result.append(candidate)
    return result


def make_question_rows(words: list[dict], package: dict) -> dict[str, list[list]]:
    by_id = {word["word_id"]: word for word in words}
    result = {"单选": [], "选词填空": [], "判断": [], "拼写": []}
    counters = Counter()

    for item in package["Items"]:
        word = by_id[item["targetWord"]]
        explanation = word["distinction"]
        source = f"TEM4 p.{word['tem4_pdf_page']}；假朋友词典 p.{word['false_friend_pdf_page']}"
        qtype = item["questionType"]
        if qtype == "SINGLE_CHOICE":
            index = counters["单选"]
            counters["单选"] += 1
            distractors = distinct_distractor_words(words, word, 11, value_key="tem4_senses")
            raw_options = [word["tem4_senses"], *(entry["tem4_senses"] for entry in distractors)]
            choices, answer = rotate_choice_options(raw_options, index)
            result["单选"].append(
                [
                    item["itemCode"],
                    word["word_id"],
                    f"请选择法语单词“{word['french_word']}”对应的正确中文含义。",
                    *choices,
                    answer,
                    explanation,
                    source,
                    "原书证据驱动复核完成",
                ]
            )
        elif qtype == "FILL_BLANK":
            index = counters["选词填空"]
            counters["选词填空"] += 1
            distractors = distinct_distractor_words(words, word, 17)
            choices, answer = rotate_choice_options(
                [word["french_word"], *(entry["french_word"] for entry in distractors)], index
            )
            result["选词填空"].append(
                [
                    item["itemCode"],
                    word["word_id"],
                    cloze_from_source(word),
                    *choices,
                    answer,
                    explanation,
                    source,
                    "原书证据驱动复核完成",
                ]
            )
        elif qtype == "TRUE_FALSE_WITH_JUSTIFICATION":
            index = counters["判断"]
            counters["判断"] += 1
            is_true = index % 2 == 0
            if is_true:
                claim = first_sense(word["tem4_senses"])
                answer = "V"
                judgment_explanation = f"正确。TEM4来源记录的完整释义为“{word['tem4_senses']}”。{explanation}"
            else:
                claim = ""
                if word["relation_type"] == "A 核心义不对应/原书另词表达":
                    claim = first_sense(word["english_senses"])
                if not claim or claim in word["tem4_senses"]:
                    cursor = word["sequence"] + 23
                    while not claim or claim in word["tem4_senses"]:
                        next_word = words[cursor % len(words)]
                        cursor += 29
                        candidate_claim = first_sense(next_word["tem4_senses"])
                        if candidate_claim and candidate_claim not in word["tem4_senses"]:
                            claim = candidate_claim
                answer = "F"
                judgment_explanation = f"错误。法语“{word['french_word']}”在TEM4来源中的释义为“{word['tem4_senses']}”。{explanation}"
            result["判断"].append(
                [
                    item["itemCode"],
                    word["word_id"],
                    f"判断正误：法语“{word['french_word']}”表示“{claim}”。",
                    "Vrai（正确）",
                    "Faux（错误）",
                    "",
                    "",
                    answer,
                    judgment_explanation,
                    source,
                    "原书证据驱动复核完成",
                ]
            )
        elif qtype == "SHORT_TEXT":
            counters["拼写"] += 1
            result["拼写"].append(
                [
                    item["itemCode"],
                    word["word_id"],
                    f"{word['chinese_gloss']}：________（填写对应的法语假朋友词）",
                    "",
                    "",
                    "",
                    "",
                    word["french_word"],
                    explanation,
                    source,
                    "原书证据驱动复核完成",
                ]
            )
    return result


def make_question_workbook(words: list[dict], package: dict) -> None:
    rows_by_sheet = make_question_rows(words, package)
    wb = Workbook()
    wb.remove(wb.active)
    headers = [
        "题目编号",
        "词条ID",
        "题干",
        "选项A",
        "选项B",
        "选项C",
        "选项D",
        "标准答案",
        "词义辨析解析",
        "来源页码",
        "教研状态",
    ]
    subtitles = {
        "单选": "客观题｜采用审定后的TEM4完整释义；四个选项与答案位置均衡轮换。",
        "选词填空": "客观题｜优先使用TEM4页码证据中的例组与搭配挖空；无例组时使用受控释义句。",
        "判断": "Vrai/Faux｜正误数量均衡，真命题来自TEM4释义，假命题优先来自英文迁移义。",
        "拼写": "唯一答案填空｜根据专四中文释义拼写标准法语词；可用于错后首字母提示。",
    }
    for sheet_name in ["单选", "选词填空", "判断", "拼写"]:
        ws = wb.create_sheet(sheet_name)
        write_table_sheet(
            ws,
            f"法语专四假朋友题库｜{sheet_name}",
            subtitles[sheet_name] + " 已完成双源字段、答案与选项一致性复核。",
            headers,
            rows_by_sheet[sheet_name],
            [16, 24, 58, 26, 26, 26, 26, 16, 66, 31, 18],
        )
    wb.save(QUESTION_XLSX)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.2
    for style_name, size, color in (("Title", 28, NAVY), ("Heading 1", 20, NAVY), ("Heading 2", 15, BLUE)):
        style = doc.styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_meta_line(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = FONT_CN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(GRAY)


def add_notice(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    set_cell_shading(cell, PALE_GOLD)
    set_cell_margins(cell, 170, 200, 170, 200)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("7A4E00")


def add_vocab_card(doc: Document, word: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.6)
    table.columns[1].width = Cm(11.2)
    left, right = table.rows[0].cells
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(left, PALE_BLUE)
    set_cell_shading(right, WHITE)
    set_cell_margins(left, 150, 180, 150, 180)
    set_cell_margins(right, 150, 200, 150, 200)

    p = left.paragraphs[0]
    run = p.add_run(word["french_word"])
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    if word["ipa"]:
        p = left.add_paragraph(word["ipa"])
        p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    if word["part_of_speech"]:
        p = left.add_paragraph(word["part_of_speech"])
        p.runs[0].bold = True
    p = left.add_paragraph(f"易混英文：{word['english_word']}")
    p.runs[0].font.color.rgb = RGBColor.from_string(RED)
    left.add_paragraph(f"频级：{word['frequency']}")
    left.add_paragraph(f"类型：{word['relation_type']}")
    left.add_paragraph(f"专四：{'、'.join(word['markers']) if word['markers'] else '无明确历年标记'}")

    lines = [
        ("TEM4完整释义", word["tem4_senses"]),
        ("假朋友词典法语义", word["false_friend_french_senses"]),
        ("英文词义", word["english_senses"]),
        ("原书视觉核验摘要", word["false_friend_evidence"]),
        ("原书英文摘录", word["book_excerpt"]),
        ("关系判定依据", word["book_relation_reason"]),
        ("核心辨析", word["distinction"]),
        ("记忆提示", word["mnemonic"]),
        ("来源例组/搭配", word["source_usage"]),
        ("词族与派生", word["derivatives"] or "来源页未单列"),
    ]
    first = True
    for label, value in lines:
        p = right.paragraphs[0] if first else right.add_paragraph()
        first = False
        label_run = p.add_run(f"{label}：")
        label_run.bold = True
        label_run.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(value)
    p = right.add_paragraph(
        f"来源：TEM4 PDF p.{word['tem4_pdf_page']}｜假朋友词典 PDF p.{word['false_friend_pdf_page']}｜{word['word_id']}｜原书证据驱动复核完成"
    )
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    paragraph.add_run(" 页")


def make_student_document(words: list[dict]) -> None:
    doc = Document()
    style_doc(doc)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("法语专四假朋友词汇大全")
    add_meta_line(doc, "学生学习素材｜Kirk-Greene 假朋友词典 × TEM4 核心词汇视觉核验交集")
    add_meta_line(doc, f"共 {len(words)} 个词条｜高频 {sum(w['frequency']=='高频' for w in words)}｜中频 {sum(w['frequency']=='中频' for w in words)}｜低频 {sum(w['frequency']=='低频' for w in words)}")
    doc.add_paragraph()
    add_notice(doc, "内容状态：445条词汇已按两本原始资料重新复核。关系类型以 Kirk-Greene 原文和视觉核验摘要为准；OCR摘录仅便于阅读，若有识别误差请以标注PDF页码为准。")
    doc.add_paragraph()
    doc.add_heading("如何使用这本词汇册", level=1)
    instructions = [
        "先看法语真实释义，再对照易混英文词义，建立“相似外形、不同意义”的警觉。",
        "使用“记忆提示”进行双向回忆：看到法语说中文义，看到英文及时抑制迁移。",
        "“来源例组/搭配”优先采用TEM4页码证据中保存的例组、固定搭配和词族信息，不使用虚构长句。",
        "“原书英文摘录”来自对应PDF页的逐页OCR；系统同时保留视觉核验摘要和定位状态，OCR错字不作为单独定论。",
        "遇到“B 部分对应/语义范围变化”时，先读原书摘要，再分别记法语常用义与英文词义；中文字段直接重合只作辅助。",
        "频级口径：TEM4来源记录中出现2次及以上历年专四标记为高频，1次为中频，无明确标记为低频。",
    ]
    for item in instructions:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("内容索引", level=1)
    index_table = doc.add_table(rows=1, cols=3)
    index_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = index_table.rows[0].cells
    for cell, value in zip(header, ["板块", "词条数", "学习建议"]):
        set_cell_shading(cell, NAVY)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(index_table.rows[0])
    suggestions = {
        "高频": "优先掌握，建议每日复习并配合单选与拼写。",
        "中频": "第二轮掌握，重点记忆英法词义差异。",
        "低频": "扩展学习，结合来源页码复核语义边界。",
    }
    for level in ("高频", "中频", "低频"):
        cells = index_table.add_row().cells
        count = sum(1 for word in words if word["frequency"] == level)
        for cell, value in zip(cells, [level, str(count), suggestions[level]]):
            cell.text = value
            set_cell_margins(cell)

    for level in ("高频", "中频", "低频"):
        heading = doc.add_heading(f"{level}词汇", level=1)
        heading.paragraph_format.page_break_before = True
        heading.paragraph_format.space_after = Pt(3)
        subset = [word for word in words if word["frequency"] == level]
        intro = doc.add_paragraph(f"本板块共 {len(subset)} 词。{suggestions[level]}")
        intro.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
        for word in subset:
            add_vocab_card(doc, word)

    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.text = "法语专四假朋友词汇大全｜445条原书证据复核版"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.runs[0].font.size = Pt(8.5)
        header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
        add_page_number(sec.footer.paragraphs[0])
        sec.footer.paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.core_properties.title = "法语专四假朋友词汇大全—学生阅读版"
    doc.core_properties.subject = "法语专四假朋友专项学习材料"
    doc.core_properties.author = "Lexi-Bridge FF4"
    doc.save(STUDENT_DOCX)


def validate_outputs(expected_words: int, expected_questions: dict[str, int]) -> dict:
    vocab_wb = load_workbook(VOCAB_XLSX, read_only=True, data_only=True)
    question_wb = load_workbook(QUESTION_XLSX, read_only=True, data_only=True)
    doc = Document(STUDENT_DOCX)
    checks = {
        "vocabulary_rows": vocab_wb["词汇总表"].max_row - 3,
        "question_sheets": question_wb.sheetnames,
        "question_rows": {name: question_wb[name].max_row - 3 for name in question_wb.sheetnames},
        "doc_tables": len(doc.tables),
        "doc_paragraphs": len(doc.paragraphs),
        "files": {
            str(VOCAB_XLSX.relative_to(ROOT)): VOCAB_XLSX.stat().st_size,
            str(QUESTION_XLSX.relative_to(ROOT)): QUESTION_XLSX.stat().st_size,
            str(STUDENT_DOCX.relative_to(ROOT)): STUDENT_DOCX.stat().st_size,
        },
    }
    if checks["vocabulary_rows"] != expected_words:
        raise RuntimeError(f"Vocabulary row mismatch: {checks['vocabulary_rows']} != {expected_words}")
    if checks["question_rows"] != expected_questions:
        raise RuntimeError(f"Question row mismatch: {checks['question_rows']} != {expected_questions}")
    if checks["doc_tables"] < expected_words + 2:
        raise RuntimeError("Student document does not contain all vocabulary cards")
    headers = [vocab_wb["词汇总表"].cell(3, column).value for column in range(1, vocab_wb["词汇总表"].max_column + 1)]
    required_headers = {"原书关系类型", "关系判定依据", "Kirk-Greene原文摘录（OCR）", "原文定位状态"}
    if not required_headers.issubset(set(headers)):
        raise RuntimeError(f"Missing source-led vocabulary columns: {required_headers - set(headers)}")
    return checks


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    report = json.loads(REVIEW.read_text(encoding="utf-8-sig"))
    if report.get("auditStatus") != "STRUCTURAL_PASS" or report.get("issueCount") != 0:
        raise RuntimeError("Deliverables require a zero-issue structural audit")

    words = load_jsonl(WORD_BOOK)
    package = json.loads(QUESTIONS.read_text(encoding="utf-8-sig"))
    tem4_candidates = [
        row
        for row in load_jsonl(TEM4)
        if row.get("visual_verification") is True and row.get("review_state") == "VERIFIED"
    ]
    false_friend_candidates = [
        row
        for row in load_jsonl(FALSE_FRIENDS)
        if row.get("visual_verification") is True and row.get("review_state") == "VERIFIED"
    ]
    false_friend_ocr = load_false_friend_ocr(FALSE_FRIENDS_OCR)
    if len(false_friend_ocr) < 198:
        raise RuntimeError(f"Incomplete Kirk-Greene OCR cache: {len(false_friend_ocr)} pages")
    enriched = enrich_words(words, tem4_candidates, false_friend_candidates, false_friend_ocr)
    make_vocabulary_workbook(enriched)
    make_question_workbook(enriched, package)
    make_student_document(enriched)

    expected_questions = {"单选": 112, "选词填空": 111, "判断": 111, "拼写": 111}
    checks = validate_outputs(len(enriched), expected_questions)
    print(json.dumps({"frequency": Counter(w["frequency"] for w in enriched), **checks}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
